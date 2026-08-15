from __future__ import annotations

from datetime import UTC, datetime
from html import escape
from pathlib import Path
from re import sub

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, HTTPException, status
from pydantic import BaseModel, Field
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

DATA_DIR = Path(__file__).resolve().parents[1] / "data"
EXPORT_DIR = DATA_DIR / "exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

router = APIRouter(prefix="/api/v1/export", tags=["export"])


class LetterExportRequest(BaseModel):
    sender_name: str = Field(default="Studio M 360 GmbH", min_length=2, max_length=150)
    sender_address: str = Field(
        default="Musterstraße 12, 10115 Berlin", min_length=5, max_length=200
    )
    recipient_name: str = Field(default="Max Mustermann", min_length=2, max_length=150)
    recipient_company: str | None = Field(default="Musterfirma AG", max_length=150)
    recipient_address: str = Field(
        default="Hauptstraße 45, 80331 München", min_length=5, max_length=200
    )
    date: str | None = Field(default=None, max_length=50)
    subject: str = Field(
        default="Ihr Angebot und Informationen zum Projekt", min_length=2, max_length=200
    )
    reference: str | None = Field(default="Ref-Nr. LumOS-2026-0815", max_length=100)
    salutation: str = Field(default="Sehr geehrte Damen und Herren,", min_length=2, max_length=100)
    body_text: str = Field(
        default="vielen Dank für Ihre Anfrage. Anbei erhalten Sie die gewünschten Informationen.",
        min_length=5,
        max_length=10000,
    )
    closing: str = Field(default="Mit freundlichen Grüßen", min_length=2, max_length=100)
    signoff_name: str = Field(default="Ihr LumOS Team", min_length=2, max_length=150)
    export_formats: list[str] = Field(default_factory=lambda: ["docx", "pdf"])
    custom_filename: str | None = Field(default=None, max_length=100)
    human_approved: bool = Field(default=False)


class LetterPreviewResponse(BaseModel):
    formatted_preview_html: str
    sender: str
    recipient: str
    subject: str
    date: str
    body_paragraphs: list[str]
    word_count: int
    character_count: int


class LetterGenerateResponse(BaseModel):
    success: bool
    docx_path: str | None
    pdf_path: str | None
    export_dir: str
    filename_base: str
    human_approved: bool
    created_at: str


def validate_export_filename(name: str | None) -> str:
    if not name:
        timestamp = datetime.now(UTC).strftime("%Y%m%d_%H%M%S")
        return f"geschaeftsbrief_{timestamp}"
    cleaned = sub(r"[^\w\-]", "_", name.strip()).strip("_")
    return cleaned or "geschaeftsbrief"


def create_letter_docx(data: LetterExportRequest, output_path: Path) -> Path:
    doc = Document()

    # Page Margins (DIN 5008 approx)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Small Sender Line
    p_sender_line = doc.add_paragraph()
    r_sender = p_sender_line.add_run(f"{data.sender_name} · {data.sender_address}")
    r_sender.font.size = Pt(8)
    r_sender.font.color.rgb = RGBColor(120, 120, 120)
    p_sender_line.paragraph_format.space_after = Pt(18)

    # Recipient Field
    p_rec = doc.add_paragraph()
    if data.recipient_company:
        p_rec.add_run(f"{data.recipient_company}\n")
    p_rec.add_run(f"{data.recipient_name}\n{data.recipient_address}")
    p_rec.paragraph_format.space_after = Pt(30)

    # Date & Reference Line
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    current_date = data.date or datetime.now(UTC).strftime("%d.%m.%Y")
    meta_text = f"Datum: {current_date}"
    if data.reference:
        meta_text = f"{data.reference}  |  {meta_text}"
    r_meta = p_meta.add_run(meta_text)
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(80, 80, 80)
    p_meta.paragraph_format.space_after = Pt(24)

    # Subject Line
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run(data.subject)
    r_sub.font.bold = True
    r_sub.font.size = Pt(13)
    p_sub.paragraph_format.space_after = Pt(18)

    # Salutation
    p_sal = doc.add_paragraph(data.salutation)
    p_sal.paragraph_format.space_after = Pt(12)

    # Body Paragraphs
    paragraphs = [p.strip() for p in data.body_text.split("\n") if p.strip()]
    for text in paragraphs:
        p_body = doc.add_paragraph(text)
        p_body.paragraph_format.space_after = Pt(10)
        p_body.paragraph_format.line_spacing = 1.15

    # Closing & Signoff
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(18)
    p_close.paragraph_format.space_after = Pt(36)
    p_close.add_run(f"{data.closing},\n\n")

    p_sig = doc.add_paragraph(data.signoff_name)
    r_sig = p_sig.runs[0] if p_sig.runs else p_sig.add_run(data.signoff_name)
    r_sig.font.bold = True

    doc.save(str(output_path))
    return output_path


def create_letter_pdf(data: LetterExportRequest, output_path: Path) -> Path:
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
    )
    styles = getSampleStyleSheet()

    sender_style = ParagraphStyle(
        "SenderLine",
        parent=styles["Normal"],
        fontSize=8,
        textColor="#666666",
        spaceAfter=14,
    )
    recipient_style = ParagraphStyle(
        "Recipient",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=24,
    )
    meta_style = ParagraphStyle(
        "MetaRight",
        parent=styles["Normal"],
        fontSize=9.5,
        alignment=2,  # Right
        textColor="#444444",
        spaceAfter=20,
    )
    subject_style = ParagraphStyle(
        "Subject",
        parent=styles["Heading2"],
        fontSize=13,
        leading=16,
        fontName="Helvetica-Bold",
        spaceAfter=16,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10.5,
        leading=15,
        spaceAfter=10,
    )

    story = []

    # Sender Line
    sender_text = escape(f"{data.sender_name} · {data.sender_address}")
    story.append(Paragraph(sender_text, sender_style))

    # Recipient
    rec_text = ""
    if data.recipient_company:
        rec_text += f"<b>{escape(data.recipient_company)}</b><br/>"
    rec_text += f"{escape(data.recipient_name)}<br/>{escape(data.recipient_address).replace('\n', '<br/>')}"
    story.append(Paragraph(rec_text, recipient_style))

    # Meta Line
    current_date = data.date or datetime.now(UTC).strftime("%d.%m.%Y")
    meta_str = f"Datum: {escape(current_date)}"
    if data.reference:
        meta_str = f"{escape(data.reference)}  |  {meta_str}"
    story.append(Paragraph(meta_str, meta_style))

    # Subject
    story.append(Paragraph(escape(data.subject), subject_style))

    # Salutation
    story.append(Paragraph(escape(data.salutation), body_style))

    # Body
    paragraphs = [p.strip() for p in data.body_text.split("\n") if p.strip()]
    for p in paragraphs:
        story.append(Paragraph(escape(p), body_style))

    story.append(Spacer(1, 15))

    # Closing & Signoff
    closing_text = f"{escape(data.closing)},<br/><br/><br/><b>{escape(data.signoff_name)}</b>"
    story.append(Paragraph(closing_text, body_style))

    doc.build(story)
    return output_path


@router.post(
    "/letter/preview",
    response_model=LetterPreviewResponse,
    name="preview_letter_export",
)
async def preview_letter_export(request: LetterExportRequest) -> LetterPreviewResponse:
    paragraphs = [p.strip() for p in request.body_text.split("\n") if p.strip()]
    current_date = request.date or datetime.now(UTC).strftime("%d.%m.%Y")

    rec_comp = f"<b>{escape(request.recipient_company)}</b><br/>" if request.recipient_company else ""
    html_preview = f"""
<div class="letter-preview-box">
  <div class="preview-sender">{escape(request.sender_name)} · {escape(request.sender_address)}</div>
  <div class="preview-recipient">{rec_comp}{escape(request.recipient_name)}<br/>{escape(request.recipient_address)}</div>
  <div class="preview-meta">{escape(request.reference or '')} · {escape(current_date)}</div>
  <div class="preview-subject">{escape(request.subject)}</div>
  <div class="preview-salutation">{escape(request.salutation)}</div>
  <div class="preview-body">
    {''.join(f'<p>{escape(p)}</p>' for p in paragraphs)}
  </div>
  <div class="preview-closing">{escape(request.closing)},<br/><br/><b>{escape(request.signoff_name)}</b></div>
</div>
""".strip()

    word_count = len(request.body_text.split())
    character_count = len(request.body_text)

    return LetterPreviewResponse(
        formatted_preview_html=html_preview,
        sender=f"{request.sender_name}, {request.sender_address}",
        recipient=f"{request.recipient_name}, {request.recipient_address}",
        subject=request.subject,
        date=current_date,
        body_paragraphs=paragraphs,
        word_count=word_count,
        character_count=character_count,
    )


@router.post(
    "/letter/generate",
    response_model=LetterGenerateResponse,
    name="generate_letter_export",
)
async def generate_letter_export(request: LetterExportRequest) -> LetterGenerateResponse:
    if not request.human_approved:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Der Dokumentexport erfordert eine ausdrückliche menschliche Freigabe (human_approved=True).",
        )

    base_name = validate_export_filename(request.custom_filename)
    docx_file_path: Path | None = None
    pdf_file_path: Path | None = None

    if "docx" in request.export_formats:
        target_docx = EXPORT_DIR / f"{base_name}.docx"
        docx_file_path = create_letter_docx(request, target_docx)

    if "pdf" in request.export_formats:
        target_pdf = EXPORT_DIR / f"{base_name}.pdf"
        pdf_file_path = create_letter_pdf(request, target_pdf)

    return LetterGenerateResponse(
        success=True,
        docx_path=str(docx_file_path) if docx_file_path else None,
        pdf_path=str(pdf_file_path) if pdf_file_path else None,
        export_dir=str(EXPORT_DIR),
        filename_base=base_name,
        human_approved=True,
        created_at=datetime.now(UTC).isoformat(),
    )
